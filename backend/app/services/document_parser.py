from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError


class UnsupportedDocumentError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


class MalformedDocumentError(ValueError):
    pass


class DocumentParser:
    supported_extensions = {".txt", ".md", ".markdown", ".pdf", ".docx"}

    def parse(self, filename: str, content: bytes) -> str:
        extension = Path(filename).suffix.lower()
        if extension not in self.supported_extensions:
            supported = ", ".join(sorted(self.supported_extensions))
            raise UnsupportedDocumentError(
                f"unsupported file type: {extension or 'unknown'}; supported: {supported}"
            )
        if not content:
            raise EmptyDocumentError("uploaded document is empty")

        if extension in {".txt", ".md", ".markdown"}:
            text = self._parse_text(content)
        elif extension == ".pdf":
            text = self._parse_pdf(content)
        else:
            text = self._parse_docx(content)

        normalized = text.strip()
        if not normalized:
            raise EmptyDocumentError("document does not contain extractable text")
        return normalized

    @staticmethod
    def _parse_text(content: bytes) -> str:
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise UnsupportedDocumentError("document encoding must be UTF-8 or GB18030")

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
        except (PdfReadError, ValueError, OSError) as exc:
            raise MalformedDocumentError("unable to parse PDF document") from exc
        return "\n\n".join(page.strip() for page in pages if page.strip())

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        try:
            document = DocxDocument(BytesIO(content))
        except (BadZipFile, PackageNotFoundError, ValueError, KeyError) as exc:
            raise MalformedDocumentError("unable to parse DOCX document") from exc

        blocks = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks)
