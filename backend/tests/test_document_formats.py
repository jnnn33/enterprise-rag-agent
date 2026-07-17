from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject


def _build_pdf(text: str) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
    page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(output)
    return output.getvalue()


def _build_docx() -> bytes:
    output = BytesIO()
    document = DocxDocument()
    document.add_heading("Expense policy", level=1)
    document.add_paragraph("The hotel limit is 800 dollars per night.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Category"
    table.cell(0, 1).text = "Limit"
    table.cell(1, 0).text = "Taxi"
    table.cell(1, 1).text = "100 dollars"
    document.save(output)
    return output.getvalue()


def test_pdf_upload_is_searchable(client: TestClient) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "remote-policy.pdf",
                _build_pdf("Remote work allowance is 300 dollars."),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 201
    assert response.json()["source"] == "upload:remote-policy.pdf"

    answer = client.post(
        "/api/v1/chat",
        json={"question": "What is the remote work allowance?"},
    )
    assert answer.status_code == 200
    assert "300 dollars" in answer.json()["answer"]


def test_docx_paragraphs_and_tables_are_searchable(client: TestClient) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "expense-policy.docx",
                _build_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201

    answer = client.post(
        "/api/v1/chat",
        json={"question": "What is the Taxi limit?"},
    )
    assert answer.status_code == 200
    assert "Taxi | 100 dollars" in answer.json()["answer"]


def test_malformed_pdf_is_rejected(client: TestClient) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("broken.pdf", b"not a pdf", "application/pdf")},
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "unable to parse PDF document"
